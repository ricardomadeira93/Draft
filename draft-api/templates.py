"""
templates.py — DOCX branded template management.

Orgs can upload a .docx file with {{Question}}, {{Answer}}, {{Sources}} placeholders.
When exporting, each placeholder row in the template is cloned and filled.
Falls back to a generic built-in document if no custom template exists.
"""

import io
import os
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


def _template_path(org_id: str) -> Path:
    # Sanitize org_id to be a safe filename
    safe_id = "".join(c for c in org_id if c.isalnum() or c in "-_")
    return TEMPLATES_DIR / f"{safe_id}.docx"


def save_template(org_id: str, file_bytes: bytes) -> None:
    """Save the org's branded .docx template to disk."""
    _template_path(org_id).write_bytes(file_bytes)


def delete_template(org_id: str) -> bool:
    """Remove a custom template. Returns True if it existed."""
    path = _template_path(org_id)
    if path.exists():
        path.unlink()
        return True
    return False


def has_template(org_id: str) -> bool:
    """Check if the org has a custom template uploaded."""
    return _template_path(org_id).exists()


def _replace_placeholders_in_paragraph(para, replacements: dict) -> None:
    """Replace {{Key}} placeholders across runs in a single paragraph."""
    full_text = "".join(r.text for r in para.runs)
    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
    # Write back — clear all runs except the first, put the full text there
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""


def _fill_template(template_bytes: bytes, rows: list[dict]) -> bytes:
    """
    Clone the template for each Q&A row, filling in placeholders.
    Appends a page break between rows (except after the last one).
    """
    doc = Document(io.BytesIO(template_bytes))
    result_doc = Document()

    # Copy styles from template to result
    result_doc.core_properties.title = "RFP Answers"

    for i, row in enumerate(rows):
        replacements = {
            "Question": row.get("Question", ""),
            "Answer": row.get("Answer", ""),
            "Sources": ", ".join(s.get("source", "") for s in row.get("Sources", [])),
        }

        # Deep-copy each paragraph from the template and fill placeholders
        for para in doc.paragraphs:
            new_para = copy.deepcopy(para)
            _replace_placeholders_in_paragraph(new_para, replacements)
            result_doc.element.body.append(new_para._element)

        # Page break between entries
        if i < len(rows) - 1:
            result_doc.add_page_break()

    buf = io.BytesIO()
    result_doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_docx(org_id: str, rows: list[dict]) -> bytes:
    """
    Export rows as a DOCX. Uses the org's branded template if available,
    otherwise falls back to a clean built-in document.
    """
    template_path = _template_path(org_id)

    if template_path.exists():
        return _fill_template(template_path.read_bytes(), rows)

    # Generic fallback
    doc = Document()
    doc.add_heading("RFP Answers", 0)
    for row in rows:
        doc.add_heading(row.get("Question", ""), level=1)
        doc.add_paragraph(row.get("Answer", ""))
        sources = ", ".join(s.get("source", "") for s in row.get("Sources", []))
        p = doc.add_paragraph()
        p.add_run(f"Sources: {sources}").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
