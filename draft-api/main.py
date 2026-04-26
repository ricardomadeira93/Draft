import os
import io
import csv
import ssl
import certifi

# Bypass macOS Python SSL certificate issues locally
os.environ["SSL_CERT_FILE"] = certifi.where()
ssl._create_default_https_context = ssl._create_unverified_context
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore

from database import index, INDEX_NAME
from llm_router import get_embedding_model
from rag_engine import answer_question_with_sources, _get_retriever, extract_sources
from tracker import (
    init_db, record_upload, list_uploads, delete_upload,
    create_share_session, get_share_session, update_share_answer,
    save_document_history, list_document_history, get_document_history_answers
)
from auth import get_current_org
from fastapi import Depends

app = FastAPI(title="Draft API", description="RAG-powered RFP automation backend")

init_db()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Health ───────────────────────────────────────────────────────────────────

@app.get("/")
def read_root():
    return {"status": "success", "message": "Draft API is running."}


# ─── Knowledge Base ───────────────────────────────────────────────────────────

@app.post("/upload-kb")
async def upload_knowledge_base(file: UploadFile = File(...), user: dict = Depends(get_current_org)):
    """
    Chunk a document, embed it, and push to Pinecone.
    Injects source filename as vector metadata for targeted deletion.
    """
    content = await file.read()
    text = content.decode("utf-8")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
    )
    chunks = text_splitter.split_text(text)
    metadatas = [{"source": file.filename} for _ in chunks]

    embeddings = get_embedding_model()
    PineconeVectorStore.from_texts(
        texts=chunks,
        embedding=embeddings,
        index_name=INDEX_NAME,
        metadatas=metadatas,
        namespace=user["org_id"]
    )

    record_upload(user["org_id"], file.filename, len(chunks))

    return {
        "status": "success",
        "message": f"Successfully indexed '{file.filename}' into {len(chunks)} vectors.",
    }


@app.get("/kb/files")
def get_kb_files(user: dict = Depends(get_current_org)):
    """List all documents tracked in the Knowledge Base."""
    return {"files": list_uploads(user["org_id"])}


@app.delete("/kb/files/{filename}")
def delete_kb_file(filename: str, user: dict = Depends(get_current_org)):
    """Purge a document's vectors from Pinecone and remove it from the tracker."""
    try:
        index.delete(filter={"source": {"$eq": filename}}, namespace=user["org_id"])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Pinecone deletion failed: {e}")

    if not delete_upload(user["org_id"], filename):
        raise HTTPException(status_code=404, detail=f"'{filename}' not found in tracker.")

    return {"status": "success", "message": f"Deleted '{filename}' from Knowledge Base."}


# ─── RFP Processing ───────────────────────────────────────────────────────────

@app.post("/process-csv")
async def process_csv(file: UploadFile = File(...), language: str = Form("English"), user: dict = Depends(get_current_org)):
    """
    Process a CSV with a 'Question' column.
    Returns JSON with answers AND source attribution per row.
    """
    content = await file.read()
    csv_reader = csv.DictReader(io.StringIO(content.decode("utf-8")))

    results = []
    for row in csv_reader:
        question = row.get("Question", "").strip()
        if not question:
            continue
        result = await answer_question_with_sources(question, language, user["org_id"])
        results.append({
            "Question": question,
            "Answer": result["answer"],
            "Sources": result["sources"],
        })

    import uuid
    session_id = str(uuid.uuid4())
    if results:
        save_document_history(session_id, user["org_id"], file.filename, len(results), results)

    return {"results": results, "session_id": session_id}

@app.get("/history")
def get_history(user: dict = Depends(get_current_org)):
    """List all past processed documents for an organization."""
    return {"history": list_document_history(user["org_id"])}

@app.get("/history/{session_id}")
def get_history_detail(session_id: str, user: dict = Depends(get_current_org)):
    """Fetch the full answers for a past processed document."""
    answers = get_document_history_answers(user["org_id"], session_id)
    if not answers:
        raise HTTPException(status_code=404, detail="History not found.")
    return {"answers": answers}


# ─── Export Processing ────────────────────────────────────────────────────────

class SourceModel(BaseModel):
    source: str
    snippet: str

class QARowModel(BaseModel):
    Question: str
    Answer: str
    Sources: List[SourceModel]

class ExportRequest(BaseModel):
    results: List[QARowModel]
    format: str

@app.post("/export")
async def export_results(body: ExportRequest, user: dict = Depends(get_current_org)):
    """Generate a PDF or DOCX file from the generated answers."""
    if not body.results:
        raise HTTPException(status_code=400, detail="No results provided to export.")
    
    if body.format == "pdf":
        try:
            from weasyprint import HTML
            # Generate a simple HTML string to convert to PDF
            html_content = "<html><head><style>body { font-family: sans-serif; }</style></head><body>"
            html_content += "<h1>RFP Answers</h1>"
            for row in body.results:
                html_content += f"<h2>Q: {row.Question}</h2>"
                html_content += f"<p><strong>A:</strong> {row.Answer}</p>"
                sources = ", ".join([s.source for s in row.Sources])
                html_content += f"<p><small><em>Sources: {sources}</em></small></p><hr>"
            html_content += "</body></html>"
            
            pdf_bytes = HTML(string=html_content).write_pdf()
            return StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=rfp_answers.pdf"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")
            
    elif body.format == "docx":
        try:
            from docx import Document
            document = Document()
            document.add_heading('RFP Answers', 0)
            
            for row in body.results:
                document.add_heading(row.Question, level=1)
                document.add_paragraph(row.Answer)
                sources = ", ".join([s.source for s in row.Sources])
                p = document.add_paragraph()
                p.add_run(f"Sources: {sources}").italic = True
                
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            
            return StreamingResponse(
                file_stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=rfp_answers.docx"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")
            
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")


# ─── Share / Review Sessions ──────────────────────────────────────────────────

import uuid
from datetime import datetime, timedelta, timezone

class ShareRequest(BaseModel):
    results: List[QARowModel]

@app.post("/share")
async def create_share(body: ShareRequest, user: dict = Depends(get_current_org)):
    """Create a shareable link for review, valid for 7 days."""
    if not body.results:
        raise HTTPException(status_code=400, detail="No results provided to share.")
        
    session_id = str(uuid.uuid4())
    token = str(uuid.uuid4()) # For public link
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    # Store in DB using org_id
    create_share_session(session_id, user["org_id"], token, [r.model_dump() for r in body.results], expires_at)
    
    # Normally we'd use an env var like SHARE_BASE_URL, but for now we return just the token.
    # The frontend will construct the full URL.
    return {
        "status": "success",
        "token": token,
        "expires_at": expires_at.isoformat()
    }

@app.get("/review/{token}")
def get_review(token: str):
    """Fetch a shared session for public review."""
    session = get_share_session(token)
    if not session:
        raise HTTPException(status_code=404, detail="Review session not found or expired.")
    
    # Only return the data needed for the review UI
    return {
        "status": session["status"],
        "answers": session["answers"]
    }

class ReviewUpdateRequest(BaseModel):
    status: str # 'approved' or 'flagged'
    comment: str = ""

@app.patch("/review/{token}/answer/{index}")
def update_review(token: str, index: int, body: ReviewUpdateRequest):
    """Update the status of a specific answer in a review session."""
    if body.status not in ["approved", "flagged"]:
        raise HTTPException(status_code=400, detail="Status must be 'approved' or 'flagged'.")
        
    success = update_share_answer(token, index, body.status, body.comment)
    if not success:
        raise HTTPException(status_code=404, detail="Review session or answer not found.")
        
    return {"status": "success"}


# ─── Evaluation Dashboard ─────────────────────────────────────────────────────

class EvaluateRequest(BaseModel):
    question: str
    language: str = "English"


@app.post("/evaluate")
async def evaluate(body: EvaluateRequest, user: dict = Depends(get_current_org)):
    """
    Single-question evaluation endpoint.
    Returns the answer, the raw retrieved chunks, and source attribution.
    Used by the Evaluation Dashboard to make the RAG pipeline transparent.
    """
    question = body.question.strip()
    if not question:
        raise HTTPException(status_code=422, detail="Question cannot be empty.")

    docs = await _get_retriever(user["org_id"]).ainvoke(question)
    result = await answer_question_with_sources(question, body.language, user["org_id"])

    chunks = [
        {
            "source": doc.metadata.get("source", "Unknown"),
            "text": doc.page_content,
            "rank": i + 1,
        }
        for i, doc in enumerate(docs)
    ]

    return {
        "question": question,
        "answer": result["answer"],
        "sources": result["sources"],
        "retrieved_chunks": chunks,
    }